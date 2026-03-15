"""
Tests for File Upload and Parsing Functionality
"""
import pytest
import io
from pathlib import Path
import sys
from docx import Document

# Add backend to path
sys.path.insert(0, str(Path(__file__).parent.parent))

from core.config import settings
from services.file_parsing_service import file_parsing_service
from services.storage_service import storage_service
from schemas.file_upload import ProgrammingLanguage


# Sample code for testing
PYTHON_CODE = '''
import os
from typing import List, Optional

def hello_world(name: str) -> str:
    """Greet someone."""
    return f"Hello, {name}!"

class Calculator:
    """A simple calculator."""
    def add(self, a: int, b: int) -> int:
        return a + b
    
    def subtract(self, a: int, b: int) -> int:
        return a - b

PI = 3.14159
'''

JAVASCRIPT_CODE = '''
import React from 'react';
import { useState } from 'react';

const greeting = "Hello";

function sayHello(name) {
    return `Hello, ${name}!`;
}

class Calculator {
    add(a, b) {
        return a + b;
    }
}

export default Calculator;
'''

JAVA_CODE = '''
import java.util.List;
import java.util.ArrayList;

public class Calculator {
    private int value;
    
    public int add(int a, int b) {
        return a + b;
    }
    
    public int subtract(int a, int b) {
        return a - b;
    }
}
'''

C_CODE = '''
#include <stdio.h>
#include <stdlib.h>

int add(int a, int b) {
    return a + b;
}

int main() {
    printf("Hello, World!");
    return 0;
}
'''

PHP_CODE = '''
<?php

use App\\Services\\Mailer;
require_once 'bootstrap.php';

$config = ['env' => 'test'];

class InvoiceService extends BaseService {
    public function sendInvoice($userId, string $email) {
        return $email;
    }
}

function helper($name, $age = 18) {
    return $name;
}
'''


class TestFileParsingService:
    """Test cases for file parsing service."""
    
    @pytest.mark.asyncio
    async def test_detect_language_python(self):
        """Test language detection for Python files."""
        lang = file_parsing_service.detect_language(".py")
        assert lang == ProgrammingLanguage.PYTHON
    
    @pytest.mark.asyncio
    async def test_detect_language_javascript(self):
        """Test language detection for JavaScript files."""
        lang = file_parsing_service.detect_language(".js")
        assert lang == ProgrammingLanguage.JAVASCRIPT
        
        lang = file_parsing_service.detect_language(".jsx")
        assert lang == ProgrammingLanguage.JSX
    
    @pytest.mark.asyncio
    async def test_detect_language_java(self):
        """Test language detection for Java files."""
        lang = file_parsing_service.detect_language(".java")
        assert lang == ProgrammingLanguage.JAVA
    
    @pytest.mark.asyncio
    async def test_detect_language_c_cpp(self):
        """Test language detection for C/C++ files."""
        assert file_parsing_service.detect_language(".c") == ProgrammingLanguage.C
        assert file_parsing_service.detect_language(".cpp") == ProgrammingLanguage.CPP
        assert file_parsing_service.detect_language(".h") == ProgrammingLanguage.C
        assert file_parsing_service.detect_language(".hpp") == ProgrammingLanguage.CPP

    @pytest.mark.asyncio
    async def test_detect_language_php(self):
        """Test language detection for PHP files."""
        assert file_parsing_service.detect_language(".php") == ProgrammingLanguage.PHP
    
    @pytest.mark.asyncio
    async def test_parse_python_file(self):
        """Test parsing Python code."""
        result = await file_parsing_service.parse_file(PYTHON_CODE, ".py", "test-py-001")
        
        assert result.language == ProgrammingLanguage.PYTHON
        assert result.syntax_validation.is_valid
        assert len(result.structure.imports) >= 2
        assert len(result.structure.functions) >= 1
        assert len(result.structure.classes) >= 1
        assert "PI" in result.structure.global_variables
        assert result.metrics.total_lines > 0
    
    @pytest.mark.asyncio
    async def test_parse_javascript_file(self):
        """Test parsing JavaScript code."""
        result = await file_parsing_service.parse_file(JAVASCRIPT_CODE, ".js", "test-js-001")
        
        assert result.language == ProgrammingLanguage.JAVASCRIPT
        assert len(result.structure.imports) >= 1
        assert len(result.structure.functions) >= 1
        assert len(result.structure.classes) >= 1
    
    @pytest.mark.asyncio
    async def test_parse_java_file(self):
        """Test parsing Java code."""
        result = await file_parsing_service.parse_file(JAVA_CODE, ".java", "test-java-001")
        
        assert result.language == ProgrammingLanguage.JAVA
        assert len(result.structure.imports) >= 2
        assert len(result.structure.classes) >= 1
    
    @pytest.mark.asyncio
    async def test_parse_c_file(self):
        """Test parsing C code."""
        result = await file_parsing_service.parse_file(C_CODE, ".c", "test-c-001")
        
        assert result.language == ProgrammingLanguage.C
        assert len(result.structure.imports) >= 2  # includes
        assert len(result.structure.functions) >= 1

    @pytest.mark.asyncio
    async def test_parse_php_file(self):
        """Test parsing PHP code."""
        result = await file_parsing_service.parse_file(PHP_CODE, ".php", "test-php-001")

        assert result.language == ProgrammingLanguage.PHP
        assert result.syntax_validation.is_valid
        assert len(result.structure.imports) >= 2
        assert len(result.structure.functions) >= 2
        assert len(result.structure.classes) >= 1
        assert "config" in result.structure.global_variables
    
    @pytest.mark.asyncio
    async def test_basic_metrics(self):
        """Test basic metrics calculation."""
        result = await file_parsing_service.parse_file(PYTHON_CODE, ".py", "test-metrics")
        
        assert result.metrics.total_lines > 0
        assert result.metrics.code_lines > 0
        assert result.metrics.function_count >= 1
        assert result.metrics.class_count >= 1
        assert result.metrics.import_count >= 2


class TestFileUploadApi:
    def _set_local_storage(self, monkeypatch, tmp_path):
        monkeypatch.setattr(settings, "UPLOAD_DIR", str(tmp_path / "uploads"))
        monkeypatch.setattr(settings, "UPLOAD_STORAGE_BACKEND", "local")

    def test_upload_python_file(self, client, monkeypatch, tmp_path):
        self._set_local_storage(monkeypatch, tmp_path)
        response = client.post(
            "/api/v1/files/upload",
            files={"file": ("solution.py", PYTHON_CODE.encode("utf-8"), "text/x-python")},
            data={"uploader_id": "student_001", "assignment_id": "assignment_001"},
        )

        assert response.status_code == 201
        body = response.json()
        assert body["file_extension"] == ".py"
        assert body["language"] == "python"

    def test_upload_php_file(self, client, monkeypatch, tmp_path):
        self._set_local_storage(monkeypatch, tmp_path)
        response = client.post(
            "/api/v1/files/upload",
            files={"file": ("lesson.php", PHP_CODE.encode("utf-8"), "application/x-httpd-php")},
        )

        assert response.status_code == 201
        file_id = response.json()["file_id"]
        assert response.json()["language"] == "php"

        parse_response = client.get(f"/api/v1/files/{file_id}/parse")
        assert parse_response.status_code == 200
        assert parse_response.json()["language"] == "php"

    def test_upload_docx_file_and_parse(self, client, monkeypatch, tmp_path):
        self._set_local_storage(monkeypatch, tmp_path)
        buffer = io.BytesIO()
        document = Document()
        document.add_paragraph("这是一个 docx 上传测试。")
        document.add_paragraph("第二行内容。")
        document.save(buffer)
        buffer.seek(0)

        upload_response = client.post(
            "/api/v1/files/upload",
            files={
                "file": (
                    "report.docx",
                    buffer.getvalue(),
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
            },
        )

        assert upload_response.status_code == 201
        file_id = upload_response.json()["file_id"]

        content_response = client.get(f"/api/v1/files/{file_id}/content")
        assert content_response.status_code == 200
        assert "docx 上传测试" in content_response.json()["content"]

        parse_response = client.get(f"/api/v1/files/{file_id}/parse")
        assert parse_response.status_code == 200
        assert parse_response.json()["file_extension"] == ".docx"
        assert "第二行内容" in parse_response.json()["content"]

    def test_upload_invalid_extension(self, client, monkeypatch, tmp_path):
        self._set_local_storage(monkeypatch, tmp_path)
        response = client.post(
            "/api/v1/files/upload",
            files={"file": ("malware.exe", b"MZ", "application/octet-stream")},
        )

        assert response.status_code == 400
        assert "not allowed" in response.json()["detail"]

    def test_upload_file_size_limit(self, client, monkeypatch, tmp_path):
        self._set_local_storage(monkeypatch, tmp_path)
        monkeypatch.setattr(settings, "MAX_UPLOAD_SIZE", 8)
        response = client.post(
            "/api/v1/files/upload",
            files={"file": ("big.py", b"print('toolarge')", "text/x-python")},
        )

        assert response.status_code == 400
        assert "maximum allowed size" in response.json()["detail"]

    def test_delete_file_removes_local_object(self, client, monkeypatch, tmp_path):
        self._set_local_storage(monkeypatch, tmp_path)
        upload_response = client.post(
            "/api/v1/files/upload",
            files={"file": ("cleanup.py", PYTHON_CODE.encode("utf-8"), "text/x-python")},
        )

        file_id = upload_response.json()["file_id"]
        uploads_dir = Path(settings.UPLOAD_DIR)
        stored_files = list(uploads_dir.glob(f"{file_id}.*"))
        assert stored_files

        delete_response = client.delete(f"/api/v1/files/{file_id}")
        assert delete_response.status_code == 200
        assert not stored_files[0].exists()


class TestStorageServiceS3:
    @pytest.mark.asyncio
    async def test_s3_roundtrip(self, monkeypatch):
        objects = {}

        class FakeBody:
            def __init__(self, data):
                self._data = data

            def read(self):
                return self._data

        class FakeS3Client:
            def put_object(self, Bucket, Key, Body, ContentType):
                objects[(Bucket, Key)] = {"body": Body, "content_type": ContentType}

            def get_object(self, Bucket, Key):
                return {"Body": FakeBody(objects[(Bucket, Key)]["body"])}

            def delete_object(self, Bucket, Key):
                objects.pop((Bucket, Key), None)

        monkeypatch.setattr(settings, "UPLOAD_STORAGE_BACKEND", "s3")
        monkeypatch.setattr(settings, "UPLOAD_S3_BUCKET", "test-bucket")
        monkeypatch.setattr(settings, "UPLOAD_S3_KEY_PREFIX", "test-prefix")
        monkeypatch.setattr(storage_service, "_get_s3_client", lambda: FakeS3Client())

        result = await storage_service.save_upload(
            content=b"print('hello from s3')",
            original_filename="cloud.py",
            file_id="file_cloud_001",
            content_type="text/x-python",
        )

        assert result.file_path == "s3://test-bucket/test-prefix/file_cloud_001.py"
        stored_bytes = await storage_service.read_bytes(result.file_path)
        assert stored_bytes == b"print('hello from s3')"

        await storage_service.delete_file(result.file_path)
        assert not objects


if __name__ == "__main__":
    pytest.main([__file__, "-v"])

